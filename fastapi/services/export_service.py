import json
import logging
import re
import uuid
from dataclasses import asdict, dataclass
from datetime import datetime, timedelta, timezone
from io import BytesIO
from json import JSONDecodeError
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from firebase_admin import storage
from google.api_core.exceptions import NotFound
from google.cloud.firestore_v1 import SERVER_TIMESTAMP
from openai import AsyncOpenAI

from services.cost_service import TokenUsage, get_cost_service
from services.firebase_service import AI_GENERIC_ERROR_MESSAGE, firebase_service
from services.user_key_service import user_key_service
from utils.config import config

logger = logging.getLogger(__name__)

# -----------------------------------------------------------------------------
# Citation parsing + deterministic resolution (ported from test1.ipynb)
# -----------------------------------------------------------------------------

YEAR_RE = re.compile(r"\b(18|19|20)\d{2}[a-z]?\b")  # 2014 / 2014a

NAME_TOKEN = r"(?:[A-ZÄÖÜ][A-Za-zÄÖÜäöüß’'\-]{1,}|[A-ZÄÖÜ]{2,})"
PARTICLE = r"(?:van|von|de|del|der|den|da|di)"
NAME = rf"(?:{PARTICLE}\s+)?{NAME_TOKEN}"

# Supports:
# - "Cepeda und Martin"
# - "Schmidt & Keller"
# - "Hopf, Schmidt, Echterhoff"
# - "Homburg et al."
# - "OECD"
AUTHORS_PATTERN = rf"{NAME}(?:\s*,\s*{NAME})*(?:\s*(?:und|&)\s*{NAME})*(?:\s+et al\.)?"

# Matches BOTH:
# - "Kuckartz, 2014"
# - "Kuckartz 2014"
AUTHOR_YEAR_RE = re.compile(
    rf"(?P<authors>{AUTHORS_PATTERN})\s*,?\s*(?P<year>(18|19|20)\d{{2}}[a-z]?)"
)
PAREN_RE = re.compile(r"\(([^()]*)\)")

LOCATOR_START_RE = re.compile(
    r"(^|[\s,;])"
    r"(S\.|p\.|pp\.|Kap\.|Abs\.|Rn\.|Rn|§|Nr\.)"
    r"(?=\s*\d|\s*[ivxlcdmIVXLCDM])"
)

SHORTHAND_MARKERS_RE = re.compile(
    r"\b(ebd\.|ibid\.|ders\.|dies\.|dieser|dieselbe|dieselben|a\.a\.O\.|op\. cit\.)\b",
    re.IGNORECASE,
)

SUSPICIOUS_LOCATION_AUTHORS = {
    "Berlin",
    "München",
    "Hamburg",
    "Köln",
    "Frankfurt",
    "Stuttgart",
    "Wien",
    "Zürich",
    "Paris",
    "London",
}

SUSPICIOUS_NONAUTHOR_TOKENS = {
    "Tagung",
    "Workshop",
    "Konferenz",
    "Sitzung",
    "Kapitel",
    "Abschnitt",
    "Tabelle",
    "Abbildung",
    "Anhang",
    "Studie",
    "Projekt",
    "Bericht",
}

CITATION_FIXUP_MODEL = "gpt-5-nano"

CITATION_FIXUP_SCHEMA: dict = {
    "type": "object",
    "properties": {
        "results": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "job_id": {"type": "string"},
                    "is_citation": {"type": "boolean"},
                    "authors": {"type": ["string", "null"]},
                    "year": {"type": ["string", "null"]},
                    "locator": {"type": ["string", "null"]},
                    "confidence": {"type": "number"},
                    "note": {"type": ["string", "null"]},
                },
                "required": [
                    "job_id",
                    "is_citation",
                    "authors",
                    "year",
                    "locator",
                    "confidence",
                    "note",
                ],
                "additionalProperties": False,
            },
        }
    },
    "required": ["results"],
    "additionalProperties": False,
}


@dataclass
class ParsedCitation:
    raw_item: str
    authors: Optional[str]
    year: Optional[str]
    locator: Optional[str]
    prefix: Optional[str]
    confidence: float
    issues: List[str]


def find_citation_groups(text: str) -> List[Dict[str, Any]]:
    """
    Returns parenthesis spans that look like citations, e.g.
      (Kuckartz, 2014; Berelson, 1952; Früh, 2011, S. 42)
      (Cepeda und Martin 2005; Homburg et al. 2017)   <- NEW supported
    """
    groups: List[Dict[str, Any]] = []
    for m in PAREN_RE.finditer(text):
        inner = m.group(1).strip()

        # must contain a year-like token
        if not YEAR_RE.search(inner):
            continue

        # must contain at least one likely author-year pattern (comma optional)
        if not AUTHOR_YEAR_RE.search(inner):
            continue

        groups.append(
            {
                "start": m.start(),
                "end": m.end(),
                "raw": m.group(0),
                "inner": inner,
            }
        )
    return groups


def split_citation_items(inner: str) -> List[str]:
    parts = [p.strip() for p in (inner or "").split(";")]
    return [p for p in parts if p]


def strip_prefix(raw_item: str) -> Tuple[str, Optional[str]]:
    s = (raw_item or "").strip()
    for pref in ("vgl.", "cf.", "see", "vgl"):
        if s.lower().startswith(pref):
            rest = s[len(pref) :].lstrip(" ,")
            return rest, pref
    return s, None


def locator_looks_valid(locator: str) -> bool:
    loc = (locator or "").strip()
    if not loc:
        return False
    if LOCATOR_START_RE.search(loc):
        return True
    if re.fullmatch(r"\d{1,4}(\s*[–-]\s*\d{1,4})?(f{1,2})?", loc):
        return True
    return False


def parse_citation_item(item: str) -> ParsedCitation:
    raw = item.strip()
    core, pref = strip_prefix(raw)

    issues: List[str] = []
    confidence = 0.0

    y = YEAR_RE.search(core)
    if not y:
        return ParsedCitation(
            raw_item=raw,
            authors=None,
            year=None,
            locator=None,
            prefix=pref or None,
            confidence=0.0,
            issues=["no_year_found"],
        )

    year = y.group(0)
    before = core[: y.start()].strip()
    after = core[y.end() :].strip()

    authors = before.rstrip(" ,;")
    locator = after.lstrip(" ,;")
    locator = locator if locator else None

    if AUTHOR_YEAR_RE.search(core):
        confidence = 0.95
    else:
        confidence = 0.65
        issues.append("year_found_but_no_author_year_pattern")

    if not authors:
        confidence = min(confidence, 0.4)
        issues.append("missing_authors")

    if locator and not locator_looks_valid(locator):
        issues.append("locator_unrecognized_format")

    if SHORTHAND_MARKERS_RE.search(core):
        issues.append("shorthand_reference_needs_context")
        confidence = min(confidence, 0.6)

    return ParsedCitation(
        raw_item=raw,
        authors=authors or None,
        year=year,
        locator=locator,
        prefix=pref or None,
        confidence=confidence,
        issues=issues,
    )


def parse_citation_group(inner: str) -> List[ParsedCitation]:
    return [parse_citation_item(it) for it in split_citation_items(inner)]


def extract_locator_from_raw(raw_item: str) -> Optional[str]:
    s = (raw_item or "").strip()

    if re.match(r"^(S\.|p\.|pp\.|Kap\.|Abs\.|Rn\.|Rn|§|Nr\.)\s*", s):
        return s

    m = re.search(r"\b(S\.|p\.|pp\.|Kap\.|Abs\.|Rn\.|Rn|§|Nr\.)\s*\d", s)
    if m:
        return s[m.start() :].strip()

    if "," in s:
        after = s.split(",", 1)[1].strip()
        return after if after else None

    return None


def resolve_group_shorthand(parsed: List[ParsedCitation]) -> List[ParsedCitation]:
    resolved: List[ParsedCitation] = []
    last_full: Optional[ParsedCitation] = None

    for c in parsed:
        c2 = ParsedCitation(**asdict(c))

        if (c2.year is None or c2.authors is None) and last_full is not None:
            loc = extract_locator_from_raw(c2.raw_item)
            if loc:
                c2.authors = last_full.authors
                c2.year = last_full.year
                c2.locator = loc
                c2.issues = [
                    iss
                    for iss in c2.issues
                    if iss not in ("no_year_found", "missing_authors")
                ]
                c2.issues.append("resolved_from_previous_in_group")
                c2.confidence = max(c2.confidence, 0.90)

        if (
            c2.authors
            and re.fullmatch(
                r"(ders\.|dies\.|ibid\.|ebd\.)", c2.authors.strip(), flags=re.IGNORECASE
            )
            and last_full is not None
        ):
            c2.authors = last_full.authors
            c2.issues = [
                iss for iss in c2.issues if iss != "shorthand_reference_needs_context"
            ]
            c2.issues.append("resolved_author_from_previous_in_group")
            c2.confidence = max(c2.confidence, 0.90)

        if c2.authors is not None and c2.year is not None:
            last_full = c2

        resolved.append(c2)

    return resolved


def extract_citations_resolved(text: str) -> List[Dict[str, Any]]:
    groups = find_citation_groups(text)
    out: List[Dict[str, Any]] = []

    for gi, g in enumerate(groups, 1):
        parsed = resolve_group_shorthand(parse_citation_group(g["inner"]))
        out.append(
            {
                "group_index": gi,
                "group_span": (g["start"], g["end"]),
                "group_raw": g["raw"],
                "group_inner": g["inner"],
                "citations": [asdict(c) for c in parsed],
            }
        )

    return out


def needs_llm_fallback(c: ParsedCitation) -> Tuple[bool, List[str]]:
    reasons: List[str] = []

    if c.year is None:
        reasons.append("no_year")
    if c.authors is None:
        reasons.append("no_authors")
    if "shorthand_reference_needs_context" in c.issues:
        reasons.append("shorthand_needs_context")
    if c.confidence < 0.85:
        reasons.append("low_confidence")

    if (
        c.authors
        and c.authors.strip() in SUSPICIOUS_LOCATION_AUTHORS
        and c.locator is None
    ):
        reasons.append("suspicious_location_author")
    if (
        c.authors
        and c.authors.strip() in SUSPICIOUS_NONAUTHOR_TOKENS
        and c.locator is None
    ):
        reasons.append("suspicious_nonauthor_token")

    return (len(reasons) > 0), reasons


def build_llm_jobs(text: str, extracted: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    jobs: List[Dict[str, Any]] = []

    for grp in extracted:
        g_start, g_end = grp["group_span"]
        ctx_left = (text or "")[max(0, g_start - 60) : g_start].replace("\n", " ")
        ctx_right = (text or "")[g_end : min(len(text or ""), g_end + 60)].replace(
            "\n", " "
        )

        last_full: Optional[ParsedCitation] = None
        for item_idx, cdict in enumerate(grp["citations"], 1):
            c = ParsedCitation(**cdict)
            need, reasons = needs_llm_fallback(c)

            if need:
                jobs.append(
                    {
                        "job_id": f"g{grp['group_index']}_i{item_idx}",
                        "group_raw": grp["group_raw"],
                        "raw_item": c.raw_item,
                        "parsed_guess": {
                            "authors": c.authors,
                            "year": c.year,
                            "locator": c.locator,
                        },
                        "fallback_reasons": reasons,
                        "previous_in_group": (
                            None
                            if last_full is None
                            else {
                                "authors": last_full.authors,
                                "year": last_full.year,
                                "locator": last_full.locator,
                            }
                        ),
                        "context_left": ctx_left[-60:],
                        "context_right": ctx_right[:60],
                    }
                )

            if c.authors is not None and c.year is not None:
                last_full = c

    return jobs


def _get(obj: Any, key: str, default=None):
    if isinstance(obj, dict):
        return obj.get(key, default)
    return getattr(obj, key, default)


def extract_text_from_response(resp) -> str:
    t = _get(resp, "output_text", None)
    if isinstance(t, str) and t.strip():
        return t

    chunks: List[str] = []
    for item in _get(resp, "output", []) or []:
        if _get(item, "type") != "message":
            continue
        for part in _get(item, "content", []) or []:
            part_type = _get(part, "type")
            if part_type in ("output_text", "text"):
                txt = _get(part, "text", "")
                if txt:
                    chunks.append(txt)
    return "".join(chunks)


async def call_llm_fixups(
    *,
    client: AsyncOpenAI,
    jobs: List[Dict[str, Any]],
    model: str,
) -> Tuple[List[Dict[str, Any]], Dict[str, int]]:
    if not jobs:
        return [], {"input_tokens": 0, "cached_input_tokens": 0, "output_tokens": 0}

    instructions = (
        "Normalize author-year citations. Return ONLY JSON matching the schema. "
        "Use previous_in_group when helpful. "
        "If NOT a citation: is_citation=false and authors/year/locator null. "
        "confidence is 0..1."
    )

    resp = await client.responses.create(
        model=model,
        instructions=instructions,
        input=[
            {"role": "user", "content": json.dumps({"jobs": jobs}, ensure_ascii=False)}
        ],
        text={
            "format": {
                "type": "json_schema",
                "name": "citation_fixups",
                "schema": CITATION_FIXUP_SCHEMA,
                "strict": True,
            }
        },
        reasoning={"effort": "low"},
        max_output_tokens=1200,
        store=False,
    )

    raw = extract_text_from_response(resp).strip()
    if not raw:
        raise RuntimeError("Model returned no parsable output text (empty).")

    try:
        data = json.loads(raw)
    except JSONDecodeError as exc:
        logger.error("Failed to parse citation fixups JSON. raw_head=%r", raw[:200])
        raise RuntimeError("Failed to parse citation fixups JSON.") from exc

    usage = getattr(resp, "usage", None)
    input_tokens = int(getattr(usage, "input_tokens", 0) or 0) if usage else 0
    output_tokens = int(getattr(usage, "output_tokens", 0) or 0) if usage else 0
    cached_input_tokens = 0

    details = getattr(usage, "input_tokens_details", None) if usage else None
    if details is not None:
        cached_input_tokens = int(getattr(details, "cached_tokens", 0) or 0)

    return (data.get("results") or []), {
        "input_tokens": input_tokens,
        "cached_input_tokens": cached_input_tokens,
        "output_tokens": output_tokens,
    }


def apply_fixups_to_extracted(
    extracted: List[Dict[str, Any]],
    fixups: List[Dict[str, Any]],
) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    fixup_map = {
        f.get("job_id"): f for f in fixups if isinstance(f, dict) and f.get("job_id")
    }
    dropped: List[Dict[str, Any]] = []

    for grp in extracted:
        for idx, cdict in enumerate(grp.get("citations") or [], 1):
            job_id = f"g{grp.get('group_index')}_i{idx}"
            if job_id not in fixup_map:
                continue

            f = fixup_map[job_id]
            if not f.get("is_citation"):
                cdict["issues"] = list(
                    set((cdict.get("issues") or []) + ["llm_not_a_citation"])
                )
                cdict["confidence"] = min(
                    float(cdict.get("confidence") or 1.0),
                    float(f.get("confidence") or 0.5),
                )
                dropped.append(
                    {
                        "job_id": job_id,
                        "group_index": grp.get("group_index"),
                        "group_raw": grp.get("group_raw"),
                        "raw_item": cdict.get("raw_item"),
                        "note": f.get("note"),
                        "confidence": f.get("confidence"),
                    }
                )
                cdict["authors"] = None
                cdict["year"] = None
                cdict["locator"] = None
                continue

            cdict["authors"] = f.get("authors")
            cdict["year"] = f.get("year")
            cdict["locator"] = f.get("locator")
            cdict["confidence"] = float(
                f.get("confidence") or float(cdict.get("confidence") or 0.0)
            )
            cdict["issues"] = list(set((cdict.get("issues") or []) + ["llm_fixed"]))
            if f.get("note"):
                cdict["issues"].append(f"llm_note:{f['note']}")

    return extracted, dropped


# -----------------------------------------------------------------------------
# DOCX helpers (ported from test1.ipynb)
# -----------------------------------------------------------------------------

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _force_superscript_run(run) -> None:
    rPr = run.get_or_add_rPr()

    vert = rPr.find(qn("w:vertAlign"))
    if vert is None:
        vert = OxmlElement("w:vertAlign")
        rPr.append(vert)
    vert.set(qn("w:val"), "superscript")

    rStyle = rPr.find(qn("w:rStyle"))
    if rStyle is None:
        rStyle = OxmlElement("w:rStyle")
        rPr.insert(0, rStyle)
    rStyle.set(qn("w:val"), "FootnoteReference")


def superscript_footnote_numbers_in_footnotes(doc: Document) -> None:
    footnotes_part = None
    for rel in doc.part.rels.values():
        if rel.reltype.endswith("/footnotes"):
            footnotes_part = rel.target_part
            break
    if footnotes_part is None:
        return

    footnotes_root = footnotes_part._element
    for footnote in footnotes_root.findall(f"{{{W_NS}}}footnote"):
        ftype = footnote.get(qn("w:type"))
        if ftype in ("separator", "continuationSeparator"):
            continue

        for r in footnote.findall(f".//{{{W_NS}}}r"):
            if r.find(qn("w:footnoteRef")) is not None:
                _force_superscript_run(r)


def normalize_docx_text(s: str) -> str:
    return (
        (s or "")
        .replace("\u2013", "-")  # en-dash
        .replace("\u2014", "-")  # em-dash
        .replace("\u2011", "-")  # non-breaking hyphen
    )


def superscript_last_footnote_reference(paragraph) -> None:
    for r in reversed(paragraph._p.r_lst):
        if r.find(qn("w:footnoteReference")) is not None:
            _force_superscript_run(r)
            return


def add_footnote_safe(paragraph, footnote_text: str) -> None:
    if not hasattr(paragraph, "add_footnote"):
        raise RuntimeError(
            "Your python-docx does not support Paragraph.add_footnote(). "
            "Install python-docx-2023 or implement OOXML footnotes manually."
        )
    paragraph.add_footnote(footnote_text)
    superscript_last_footnote_reference(paragraph)


def format_footnote_short(cdict: dict) -> str:
    authors = str(cdict.get("authors") or "").strip()
    year = str(cdict.get("year") or "").strip()
    locator = str(cdict.get("locator") or "").strip()
    if locator:
        return normalize_docx_text(f"{authors}, {year}, {locator}")
    return normalize_docx_text(f"{authors}, {year}")


async def paragraph_to_docx_with_footnotes(
    *,
    doc_paragraph,
    para_text: str,
    use_llm: bool,
    openai_client: Optional[AsyncOpenAI],
    usage_acc: Dict[str, int],
    model: str,
) -> None:
    para_text = (para_text or "").strip()
    if not para_text:
        return

    extracted = extract_citations_resolved(para_text)

    if use_llm and openai_client is not None:
        jobs = build_llm_jobs(para_text, extracted)
        if jobs:
            fixups, usage = await call_llm_fixups(
                client=openai_client, jobs=jobs, model=model
            )
            usage_acc["input_tokens"] += int(usage.get("input_tokens") or 0)
            usage_acc["cached_input_tokens"] += int(
                usage.get("cached_input_tokens") or 0
            )
            usage_acc["output_tokens"] += int(usage.get("output_tokens") or 0)
            extracted, _dropped = apply_fixups_to_extracted(extracted, fixups)

    punct_re = re.compile(r"^([.,;:!?])(\s*)")

    cursor = 0
    for grp in extracted:
        start, end = grp["group_span"]
        group_raw = grp["group_raw"]

        before = para_text[cursor:start]
        if before:
            if before.endswith(" "):
                before = before[:-1]
            doc_paragraph.add_run(normalize_docx_text(before))

        items = grp["citations"]
        all_items_are_citations = all(
            (c.get("authors") and c.get("year")) for c in items
        )

        if not all_items_are_citations:
            if before and not before.endswith((" ", "\n")):
                doc_paragraph.add_run(" ")
            doc_paragraph.add_run(normalize_docx_text(group_raw))
            cursor = end
            continue

        after_slice = para_text[end:]
        match = punct_re.match(after_slice)
        consumed_len = 0
        trailing_ws = ""
        if match:
            punct = match.group(1)
            trailing_ws = match.group(2) or ""
            consumed_len = len(punct) + len(trailing_ws)
            doc_paragraph.add_run(punct)

        for j, c in enumerate(items):
            add_footnote_safe(doc_paragraph, format_footnote_short(c))
            if j < len(items) - 1:
                doc_paragraph.add_run(" ")

        if trailing_ws:
            doc_paragraph.add_run(trailing_ws)

        cursor = end + consumed_len

    after = para_text[cursor:]
    if after:
        doc_paragraph.add_run(normalize_docx_text(after))


def _split_paragraphs(text: str) -> List[str]:
    parts = re.split(r"\n\s*\n", text or "")
    return [p.strip() for p in parts if p and p.strip()]


def _sort_key_nummer(nummer: str) -> List[int]:
    out: List[int] = []
    for p in (nummer or "").split("."):
        try:
            out.append(int(p.strip()))
        except Exception:
            out.append(0)
    return out


def _sanitize_filename(name: str) -> str:
    safe = re.sub(r"[^a-zA-Z0-9._ -]", "_", (name or "").strip())
    safe = re.sub(r"\s+", " ", safe).strip()
    return safe[:120] if safe else "Export"


def _candidate_bucket_names(project_id: str, configured: str) -> List[str]:
    names: List[str] = []
    configured = str(configured or "").strip()
    if configured:
        names.append(configured)

    project_id = str(project_id or "").strip()
    if project_id:
        names.extend(
            [
                f"{project_id}.firebasestorage.app",
                f"{project_id}.appspot.com",
            ]
        )

    seen = set()
    out: List[str] = []
    for name in names:
        if not name or name in seen:
            continue
        seen.add(name)
        out.append(name)
    return out


class ExportService:
    def __init__(self):
        self.firebase = firebase_service

    async def create_export_job(
        self,
        *,
        user_id: str,
        projekt_id: str,
        selection: str,
        kapitel_ids: List[str],
    ) -> str:
        export_id = str(uuid.uuid4())
        unique_ids = [
            k for k in dict.fromkeys([x for x in kapitel_ids if (x or "").strip()])
        ]
        expires_at = datetime.now(timezone.utc) + timedelta(days=7)

        project = await self.firebase.get_project(user_id, projekt_id)
        projekt_snapshot = None
        if project:
            projekt_snapshot = {
                "id": projekt_id,
                "name": project.get("name"),
                "archived": bool(project.get("archived", False)),
            }

        doc_ref = (
            self.firebase.db.collection("users")
            .document(user_id)
            .collection("exports")
            .document(export_id)
        )

        doc_ref.set(
            {
                "exportId": export_id,
                "projektId": projekt_id,
                "projektSnapshot": projekt_snapshot,
                "selection": {
                    "type": selection,
                    "kapitelIds": unique_ids,
                    "kapitelCount": len(unique_ids),
                },
                "status": "running",
                "errorMessage": None,
                "createdAt": SERVER_TIMESTAMP,
                "startedAt": SERVER_TIMESTAMP,
                "updatedAt": SERVER_TIMESTAMP,
                "finishedAt": None,
                "expiresAt": expires_at,
                "file": None,
                "model": CITATION_FIXUP_MODEL,
                "keySource": None,
                "usage": {
                    "inputTokens": 0,
                    "cachedInputTokens": 0,
                    "outputTokens": 0,
                    "totalTokens": 0,
                },
                "costUsd": 0.0,
            },
            merge=True,
        )

        return export_id

    async def process_export_job(self, *, user_id: str, export_id: str) -> None:
        export_ref = (
            self.firebase.db.collection("users")
            .document(user_id)
            .collection("exports")
            .document(export_id)
        )

        try:
            snap = export_ref.get()
            if not snap.exists:
                raise RuntimeError("Export job not found.")
            export_doc = snap.to_dict() or {}

            selection = (
                export_doc.get("selection")
                if isinstance(export_doc.get("selection"), dict)
                else {}
            )
            kapitel_ids = (
                selection.get("kapitelIds")
                if isinstance(selection.get("kapitelIds"), list)
                else []
            )
            projekt_id = str(export_doc.get("projektId") or "").strip()

            api_key, key_source = await user_key_service.resolve_api_key_for_user(
                user_id
            )
            openai_client = AsyncOpenAI(api_key=api_key)

            chapters: List[Dict[str, Any]] = []
            for kapitel_id in kapitel_ids:
                kapitel_id = str(kapitel_id or "").strip()
                if not kapitel_id:
                    continue

                kapitel = await self.firebase.get_kapitel(user_id, kapitel_id)
                if not kapitel:
                    continue

                nummer = str(kapitel.get("nummer") or "").strip() or "?"
                title = str(kapitel.get("title") or "").strip() or "Untitled"
                depth = max(1, min(3, nummer.count(".") + 1))

                run_id = str(kapitel.get("activeRunId") or "").strip()
                latest = (
                    kapitel.get("latestRun")
                    if isinstance(kapitel.get("latestRun"), dict)
                    else {}
                )
                if not run_id:
                    run_id = str(latest.get("runId") or "").strip()

                if not run_id:
                    runs = await self.firebase.get_kapitel_runs(user_id, kapitel_id)
                    best: Optional[dict] = None
                    for r in runs or []:
                        if not isinstance(r, dict):
                            continue
                        if bool(r.get("archived", False)) is True:
                            continue
                        idx = r.get("index")
                        if not isinstance(idx, int):
                            continue
                        if best is None or idx > best.get("index", -1):
                            best = r
                    if best and best.get("id"):
                        run_id = str(best.get("id")).strip()

                if not run_id:
                    continue

                lesefluss = await self.firebase.get_lesefluss_result(
                    user_id, kapitel_id, run_id
                )
                content = (
                    (lesefluss or {}).get("content")
                    if isinstance(lesefluss, dict)
                    else None
                )
                status = (
                    str((lesefluss or {}).get("status") or "").strip()
                    if isinstance(lesefluss, dict)
                    else ""
                )

                if not isinstance(content, str) or not content.strip():
                    continue
                if status and status != "success":
                    continue

                chapters.append(
                    {
                        "kapitelId": kapitel_id,
                        "runId": run_id,
                        "nummer": nummer,
                        "title": title,
                        "depth": depth,
                        "text": content,
                    }
                )

            if not chapters:
                raise RuntimeError("Keine Kapitel mit verbessertem Text gefunden.")

            chapters.sort(key=lambda c: _sort_key_nummer(str(c.get("nummer") or "")))

            project_name = None
            projekt_snapshot = (
                export_doc.get("projektSnapshot")
                if isinstance(export_doc.get("projektSnapshot"), dict)
                else None
            )
            if projekt_snapshot:
                project_name = projekt_snapshot.get("name")
            if not project_name and projekt_id:
                project = await self.firebase.get_project(user_id, projekt_id)
                if project:
                    project_name = project.get("name")
            project_name = str(project_name or "Projekt").strip() or "Projekt"

            doc = Document()
            style = doc.styles["Normal"]
            style.font.name = "Calibri"
            style.font.size = Pt(11)

            usage_acc = {
                "input_tokens": 0,
                "cached_input_tokens": 0,
                "output_tokens": 0,
            }
            wrote_any = False

            for ch in chapters:
                depth = int(ch["depth"])
                if wrote_any and depth == 1:
                    doc.add_page_break()

                doc.add_heading(f"{ch['nummer']} {ch['title']}".strip(), level=depth)

                for para in _split_paragraphs(str(ch["text"])):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    await paragraph_to_docx_with_footnotes(
                        doc_paragraph=p,
                        para_text=para,
                        use_llm=True,
                        openai_client=openai_client,
                        usage_acc=usage_acc,
                        model=CITATION_FIXUP_MODEL,
                    )

                wrote_any = True

            superscript_footnote_numbers_in_footnotes(doc)

            buf = BytesIO()
            doc.save(buf)
            docx_bytes = buf.getvalue()

            now = datetime.now(timezone.utc)
            safe_project = _sanitize_filename(project_name)
            filename = f"{safe_project} Export {now.strftime('%Y-%m-%d %H-%M')}.docx"
            storage_path = f"users/{user_id}/exports/{export_id}/{filename}"

            upload_exc: Exception | None = None
            download_token = str(uuid.uuid4())
            content_disposition = f'attachment; filename="{filename}"'
            for bucket_name in _candidate_bucket_names(
                config.FIREBASE_PROJECT_ID, config.FIREBASE_STORAGE_BUCKET
            ):
                try:
                    bucket = storage.bucket(bucket_name)
                    blob = bucket.blob(storage_path)
                    blob.metadata = {"firebaseStorageDownloadTokens": download_token}
                    blob.content_disposition = content_disposition
                    blob.upload_from_string(
                        docx_bytes,
                        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                    upload_exc = None
                    break
                except NotFound as exc:
                    upload_exc = exc
                    continue

            if upload_exc is not None:
                raise upload_exc

            input_tokens = int(usage_acc["input_tokens"])
            cached_input_tokens = int(usage_acc["cached_input_tokens"])
            output_tokens = int(usage_acc["output_tokens"])
            usage_obj = TokenUsage.from_any(
                input_tokens, cached_input_tokens, output_tokens
            )

            cost_usd = 0.0
            if usage_obj.total_tokens > 0:
                cost_service = get_cost_service(firebase_service)
                cost_breakdown, matched_model, pricing, _match_type = (
                    await cost_service.calculate_cost(
                        model=CITATION_FIXUP_MODEL,
                        usage=usage_obj,
                    )
                )

                await cost_service.log_operation(
                    operation_type="export_docx",
                    user_id=user_id,
                    user_action_id=export_id,
                    operation_details={
                        "exportId": export_id,
                        "kapitelCount": len(chapters),
                        "selectionType": str(selection.get("type") or "selected"),
                    },
                    model=CITATION_FIXUP_MODEL,
                    usage=usage_obj,
                    cost_breakdown=cost_breakdown,
                    matched_model_key=matched_model,
                    pricing=pricing,
                    key_source=key_source,
                    projekt_id=projekt_id or None,
                    projekt_snapshot=(
                        projekt_snapshot if isinstance(projekt_snapshot, dict) else None
                    ),
                    status="success",
                )
                cost_usd = float(cost_breakdown.total_cost_usd)

            export_ref.set(
                {
                    "status": "success",
                    "errorMessage": None,
                    "updatedAt": SERVER_TIMESTAMP,
                    "finishedAt": SERVER_TIMESTAMP,
                    "keySource": key_source,
                    "file": {
                        "storagePath": storage_path,
                        "fileName": filename,
                        "contentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        "sizeBytes": int(len(docx_bytes)),
                    },
                    "usage": {
                        "inputTokens": int(input_tokens),
                        "cachedInputTokens": int(cached_input_tokens),
                        "outputTokens": int(output_tokens),
                        "totalTokens": int(usage_obj.total_tokens),
                    },
                    "costUsd": float(cost_usd),
                    "model": CITATION_FIXUP_MODEL,
                },
                merge=True,
            )
        except Exception as exc:
            logger.error(
                "Export job failed (export_id=%s): %s", export_id, exc, exc_info=True
            )
            export_ref.set(
                {
                    "status": "error",
                    "errorMessage": AI_GENERIC_ERROR_MESSAGE,
                    "updatedAt": SERVER_TIMESTAMP,
                    "finishedAt": SERVER_TIMESTAMP,
                },
                merge=True,
            )


export_service = ExportService()
